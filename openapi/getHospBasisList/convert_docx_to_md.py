import docx
import os

def convert_docx_to_markdown(docx_path, output_md_path):
    """
    docx 파일을 마크다운 형식으로 변환합니다.
    """
    doc = docx.Document(docx_path)
    
    markdown_content = []
    
    # 문서의 모든 단락 처리
    for para in doc.paragraphs:
        text = para.text.strip()
        
        if not text:
            markdown_content.append("")
            continue
        
        # 스타일에 따라 마크다운 형식 적용
        style_name = para.style.name.lower()
        
        if 'heading 1' in style_name or 'title' in style_name:
            markdown_content.append(f"# {text}")
        elif 'heading 2' in style_name:
            markdown_content.append(f"## {text}")
        elif 'heading 3' in style_name:
            markdown_content.append(f"### {text}")
        elif 'heading 4' in style_name:
            markdown_content.append(f"#### {text}")
        elif 'heading 5' in style_name:
            markdown_content.append(f"##### {text}")
        elif 'heading 6' in style_name:
            markdown_content.append(f"###### {text}")
        else:
            # 일반 텍스트
            # 볼드, 이탤릭 등의 서식 처리
            formatted_text = ""
            for run in para.runs:
                run_text = run.text
                if run.bold and run.italic:
                    run_text = f"***{run_text}***"
                elif run.bold:
                    run_text = f"**{run_text}**"
                elif run.italic:
                    run_text = f"*{run_text}*"
                formatted_text += run_text
            
            markdown_content.append(formatted_text if formatted_text else text)
    
    # 테이블 처리
    for table in doc.tables:
        markdown_content.append("")
        
        # 테이블 헤더
        header_cells = table.rows[0].cells
        header = "| " + " | ".join([cell.text.strip() for cell in header_cells]) + " |"
        markdown_content.append(header)
        
        # 구분선
        separator = "| " + " | ".join(["---" for _ in header_cells]) + " |"
        markdown_content.append(separator)
        
        # 테이블 데이터
        for row in table.rows[1:]:
            cells = row.cells
            row_text = "| " + " | ".join([cell.text.strip() for cell in cells]) + " |"
            markdown_content.append(row_text)
        
        markdown_content.append("")
    
    # 마크다운 파일로 저장
    with open(output_md_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(markdown_content))
    
    print(f"변환 완료: {output_md_path}")
    print(f"총 {len(markdown_content)} 줄 생성됨")

if __name__ == "__main__":
    docx_file = r"docs\OpenAPI활용가이드_건강보험심사평가원(병원정보서비스)_210616.docx"
    output_file = r"docs\OpenAPI활용가이드_건강보험심사평가원(병원정보서비스)_210616.md"
    
    if os.path.exists(docx_file):
        convert_docx_to_markdown(docx_file, output_file)
    else:
        print(f"파일을 찾을 수 없습니다: {docx_file}")
